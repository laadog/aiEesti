from docx import Document
import os
from LanguageModel import LanguageModel

class Section:
    def __init__(self, filename: str, content: str, title: str):
        self.filename = filename
        self.content = content
        self.title = title

class Database:
    def __init__(self, data_folder: str, language_model: LanguageModel):
        self.sections = list()
        self.folder = data_folder
        self.language_model = language_model
        self.files = self._get_files()
        self.sections = list()
        self.__populate_database()

    def _get_files(self):
        return [os.path.join(self.folder, f) for f in os.listdir(self.folder) if f.endswith(".docx")]

    def __populate_database(self):
        for file in self.files:
            self.__parseFile__(file)

    def __parseFile__(self, file: str):
        doc = Document(file)
        filename = os.path.basename(file)
        current_title = ""
        current_content = ""

        for paragraph in doc.paragraphs:
            text = paragraph.text
            if not text:
                continue

            if Database._is_title(paragraph):
                if current_content != "":
                    self.sections.append(Section(filename, current_content, current_title))
                current_title = paragraph.text
                current_content = ""

            current_content += text + "\n"

        if current_content != "":
            self.sections.append(Section(filename, current_content, current_title))
    
    def _is_title(paragraph):
        if paragraph.style.name.startswith("Heading"):
            return True
        for run in paragraph.runs:
            if run.bold or run.font.size and run.font.size.pt > 12:
                return True
        return False

    def update(self, input_file):
        doc = Document(input_file)
        updates = []

        file_name = os.path.basename(input_file)
        title = ""
        previous = list()

        for paragraph in doc.paragraphs:
            current_update = None
            text = paragraph.text
            if not text:
                continue

            if Database._is_title(paragraph):
                if len(previous) != 0:
                    original_text = "\n".join(x.text for x in previous)
                    print(title + "\n" + original_text + "\n")
                    current_text  = original_text
                    current_file  = file_name
                    current_update = ""

                    for replacement in self.sections:
                        if self.language_model.determine_change(file_name, title + "\n" +original_text, title + "\n" + current_text, current_file, replacement.filename, replacement.title + replacement.content):
                            current_text = replacement.content
                            current_file = replacement.filename
                            current_update = f"Changed section '{title}' to '{replacement.title}' from {replacement.filename}"


                    if current_text != original_text:
                        for p in previous:
                            p.text = ""
                        previous[0].text = current_text
                        updates.append(current_update)
                title = text
                previous.clear()
            else:
                previous.append(paragraph)


        if len(previous) != 0:
                original_text = "\n".join(x.text for x in previous)
                current_text  = original_text
                current_file  = file_name
                current_update = ""

                for replacement in self.sections:
                    if self.language_model.determine_change(file_name, original_text, current_text, current_file, replacement.filename, replacement.content):
                        current_text = replacement.content
                        current_file = replacement.filename
                        current_update = f"Changed section '{title}' to '{replacement.title}' from {replacement.filename}"

                if current_text != original_text:
                    for p in previous:
                        p.text = ""
                    previous[0].text = current_text
                    updates.append(current_update)





        updated_filename = f"UPDATED_{input_file}"
        doc.save(updated_filename)

        changelog_filename = f"CHANGELOG_{input_file}.txt"
        with open(changelog_filename, "w") as changelog:
            changelog.write("\n".join(updates))

        print(f"Update complete: {updated_filename}")
        print(f"Changelog saved: {changelog_filename}")

